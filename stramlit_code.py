import openai
import streamlit as st
from docx import Document
from docx import shared
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

openai.api_key = "9955a46f8fe34f92bba64fc71096758c"
openai.azure_endpoint = "https://firstopenai08.openai.azure.com/"
openai.api_type = "azure"
openai.api_version = "2023-05-15"


def ChangeWidgetFontSize(wgt_txt, wch_font_size="12px"):
    htmlstr = (
        """<script>var elements = window.parent.document.querySelectorAll('*'), i;
                    for (i = 0; i < elements.length; ++i) { if (elements[i].innerText == |wgt_txt|) 
                        { elements[i].style.fontSize='"""
        + wch_font_size
        + """';} } </script>  """
    )

    htmlstr = htmlstr.replace("|wgt_txt|", "'" + wgt_txt + "'")
    st.components.v1.html(f"{htmlstr}", height=0, width=0)


st.title("고소미")

st.markdown(
    """
<style>
.mid-font {
    font-size:20px !important;
}
</style>
""",
    unsafe_allow_html=True,
)

st.markdown(
    '<p class="mid-font">안녕하세요. 법률 자동화 서비스 고소미입니다.</p>',
    unsafe_allow_html=True,
)


typeList = ["욕설", "중고거래 사기", "성회롱/성추행", "형사합의", "스토킹"]
type = st.selectbox("무슨 속상하신 일이 있으셨나요?", typeList)
# ChangeWidgetFontSize('무슨 속상하신 일이 있으셨나요?', '20px')

st.markdown(
    '<p class="mid-font">소장 작성을 위해 정보를 제공해주세요.</p>',
    unsafe_allow_html=True,
)

with st.form("complaint"):
    st.write("고소인의 기본 정보를 수집하겠습니다.")
    isYou = st.radio(label="피해자가 본인이신가요?", options=["예", "아니요"])
    # ChangeWidgetFontSize('피해자가 본인이신가요?', '15px')
    com_name = st.text_input("성명")
    com_num = st.text_input("주민등록번호")
    com_address = st.text_input("주소")
    com_phone = st.text_input("전화번호")
    com_submitted = st.form_submit_button("고소인 정보 작성 완료")

knowAccuser = st.radio(label="피고소인 정보를 알고계신가요?", options=["아니요", "예"])
# st.write('<style>div.row-widget.stRadio > div{flex-direction:row;}</style>', unsafe_allow_html=True)
# st.markdown(
#     """<style> div[class*="stRadio"] > label > div[data-testid="stMarkdownContainer"] > p {
#     font-size: 20px;} </style>
#     """, unsafe_allow_html=True)
if knowAccuser == "예":
    with st.form("accuser"):
        st.write("피고소인의 기본 정보를 수집하겠습니다.")
        acc_name = st.text_input("성명")
        acc_num = st.text_input("주민등록번호")
        acc_address = st.text_input("주소")
        acc_phone = st.text_input("전화번호")
        acc_submitted = st.form_submit_button("피고소인 정보 작성 완료")

with st.form("acc_info"):
    st.write("사건 정보를 수집하겠습니다.")
    acc_date = st.date_input("언제 발생하셨나요?")
    st.markdown(
        """<style> div[class*="stWidgetLabel"] > label > div[data-testid="stMarkdownContainer"] > p {
        font-size: 20px;} </style>
        """,
        unsafe_allow_html=True,
    )
    content = st.text_area("고소하려는 육하원칙으로 사건에 대해 간단히 이야기해주세요.")
    acc_info_submitted = st.form_submit_button("사건 정보 작성 완료")

st.divider()

if acc_info_submitted:
    st.markdown(
        '<p class="mid-font">소장 작성을 시작하겠습니다.</p>', unsafe_allow_html=True
    )
    with st.spinner("Please Wait..."):
        # '''문서 기다리기'''
        result = openai.chat.completions.create(
            model="gpt-35-turbo-001",
            temperature=1,  # 창의적으로 답변하도록 최대치인 1로 수정
            messages=[
                {
                    "role": "assistant",
                    "content": """피고소인을 (죄목) 혐의로 고소합니다. 
                            고소인은 (일시)에 (범죄 발생지)에서 고소인의 험담을 하였습니다. 
                            이에 고소장을 제출하니 철저히 수사하여 엄벌에 처해 주시기를 바랍니다.""",
                },
                {
                    "role": "system",
                    "content": "You are a lawyer drafting a complaint in korean.",
                },
                {
                    "role": "user",
                    "content": "고소장의 사건 발생일자는 " + str(acc_date) + "이다.",
                },
                {"role": "user", "content": "고소장의 내용은 " + content + "이다."},
                {
                    "role": "user",
                    "content": "이 내용을 법률적인 어체로 고소장의 범죄 내용만 상세하고 육하원칙으로 작성해줘",
                },
            ],
        )
        filename = "소장_{%s}.doc".format(datetime.now().strftime("%m/%d/%Y, %H:%M:%S"))
        document = Document()
        styles = document.styles
        head1Font = styles["Heading 1"].font
        head1Font.size = shared.Pt(26)
        head1Font.name = "바탕체"
        head1Font.color.rgb = shared.RGBColor(0, 0, 0)
        headingP = document.add_heading("고  소  장", level=1)
        headingP.alignment = 1  # 0 or left, 1 for center, 2 right, 3 justify ....
        table = document.add_table(rows=8, cols=5)
        table.style = document.styles["Table Grid"]
        table.cell(0, 0).merge(table.cell(1, 0))
        table.cell(2, 0).merge(table.cell(3, 0))
        table.cell(4, 1).merge(table.cell(4, 4))
        table.cell(5, 0).merge(table.cell(5, 4))
        table.cell(6, 1).merge(table.cell(6, 4))
        table.cell(7, 0).merge(table.cell(7, 4))
        hdr_cells = table.rows[0].cells
        hdr_cells[0].paragraphs[0].add_run("고소인")
        hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
        table.rows[0].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[1].paragraphs[0].add_run("성명")
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[0].cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[2].paragraphs[0].add_run("홍길동")
        hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[0].cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[3].paragraphs[0].add_run("주민등록번호")
        hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[0].cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[4].paragraphs[0].add_run("0000-0000")
        hdr_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[0].cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells = table.rows[1].cells
        hdr_cells[1].paragraphs[0].add_run("주소")
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[1].cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[2].paragraphs[0].add_run("00시 00구 000")
        hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[1].cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[3].paragraphs[0].add_run("전화번호")
        hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[1].cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[4].paragraphs[0].add_run("00-00")
        hdr_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[1].cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        hdr_cells = table.rows[2].cells
        hdr_cells[0].paragraphs[0].add_run("피고소인")
        hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
        table.rows[2].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[1].paragraphs[0].add_run("성명")
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[2].cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[2].paragraphs[0].add_run("")
        hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[2].cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[3].paragraphs[0].add_run("주민등록번호")
        hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[2].cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[4].paragraphs[0].add_run("모를 경우 비워둠")
        hdr_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[2].cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells = table.rows[3].cells
        hdr_cells[1].paragraphs[0].add_run("주소")
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[3].cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[2].paragraphs[0].add_run("")
        hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[3].cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[3].paragraphs[0].add_run("전화번호")
        hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[3].cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[4].paragraphs[0].add_run("")
        hdr_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[3].cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        hdr_cells = table.rows[4].cells
        hdr_cells[0].paragraphs[0].add_run("죄명")
        hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
        table.rows[4].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[1].paragraphs[0].add_run("모욕")
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table.rows[4].cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        table.rows[5].height = shared.Cm(10)
        hdr_cells = table.rows[5].cells
        hdr_cells[0].paragraphs[0].add_run("sdfjklasfd")
        hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table.rows[5].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP

        hdr_cells = table.rows[6].cells
        hdr_cells[0].paragraphs[0].add_run("입증자료")
        hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
        table.rows[6].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[1].paragraphs[0].add_run("")
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table.rows[6].cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        table.rows[7].height = shared.Cm(3)
        hdr_cells = table.rows[7].cells
        list1 = [
            "this is the first line.\n",
            "this is the second line.\n",
            "this is the third line.\n",
        ]
        table.cell(7, 0).text = " ".join(list1)
        hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[7].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

        document.save(filename)

    st.success("!생성형 AI로 소장을 작성 완료하였습니다!")
    st.write(result.choices[0].message.content)
